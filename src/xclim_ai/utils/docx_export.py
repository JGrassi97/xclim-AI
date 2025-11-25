from __future__ import annotations

from pathlib import Path
from typing import Iterable, List, Optional, Tuple
import re

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    # Nicer spacing
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_kv(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value)
    # Justify key-value block for consistency
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)


def _is_table_sep(line: str) -> bool:
    """Return True if line is a markdown table separator (---, :---:, ---:, :---)."""
    raw = line.strip()
    if not raw or '|' not in raw:
        return False
    cells = [c.strip() for c in raw.strip('|').split('|')]
    # Remove leading/trailing empty cells introduced by pipes like | a | b |
    while cells and cells[0] == '':
        cells.pop(0)
    while cells and cells[-1] == '':
        cells.pop()
    if not cells:
        return False
    valid = True
    for c in cells:
        if not re.fullmatch(r':?-{3,}:?', c):
            valid = False
            break
    return valid


def _split_table_cells(row: str) -> List[str]:
    parts = [c.strip() for c in row.strip().strip('|').split('|')]
    # Drop empty leading/trailing artifacts
    if parts and parts[0] == '':
        parts = parts[1:]
    if parts and parts[-1] == '':
        parts = parts[:-1]
    return parts


def _detect_table(lines: List[str], idx: int) -> Tuple[Optional[dict], int]:
    """Detect a markdown table starting at lines[idx]."""
    if idx >= len(lines):
        return None, idx
    header_line = lines[idx]
    if '|' not in header_line:
        return None, idx
    if idx + 1 >= len(lines):
        return None, idx
    sep_line = lines[idx + 1]
    if not _is_table_sep(sep_line):
        return None, idx
    headers = _split_table_cells(header_line)
    aligns_raw = _split_table_cells(sep_line)
    # Filter out accidental heading rows like '# | something'
    if any(h.startswith('#') for h in headers):
        return None, idx
    aligns: List[str] = []
    for a in aligns_raw:
        if a.startswith(':') and a.endswith(':'):
            aligns.append('center')
        elif a.endswith(':'):
            aligns.append('right')
        elif a.startswith(':'):
            aligns.append('left')
        else:
            aligns.append('left')
    rows: List[List[str]] = []
    i = idx + 2
    while i < len(lines):
        raw = lines[i]
        if not raw.strip():
            break
        if '|' not in raw:
            break
        row_cells = _split_table_cells(raw)
        if not row_cells:
            break
        if len(row_cells) < len(headers):
            row_cells += [''] * (len(headers) - len(row_cells))
        elif len(row_cells) > len(headers):
            row_cells = row_cells[:len(headers)]
        rows.append(row_cells)
        i += 1
    if not headers or not rows:
        return None, idx
    spec = {
        'headers': headers,
        'aligns': aligns[:len(headers)],
        'rows': rows,
    }
    return spec, i


def _apply_table_alignment(cell_paragraph, align: str):
    if align == 'center':
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _render_table(doc: Document, spec: dict):
    headers: List[str] = spec['headers']
    rows: List[List[str]] = spec['rows']
    aligns: List[str] = spec['aligns']
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    # Style preference order
    style_candidates = ["Table Grid", "Light Grid", "Medium Grid 1 Accent 1"]
    existing = {s.name for s in doc.styles}
    for sc in style_candidates:
        if sc in existing:
            table.style = sc
            break
    table.autofit = True
    # Set uniform column widths (approx, Word may adjust)
    try:
        total_width = Inches(6.0)
        col_width = total_width / ncols
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
    except Exception:
        pass
    # Header
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        p = hdr_cells[j].paragraphs[0]
        p.clear()
        _append_markdown_inline(p, h)
        for run in p.runs:
            run.bold = True
        _apply_table_alignment(p, aligns[j] if j < len(aligns) else 'left')
    # Body with zebra striping
    for i, r in enumerate(rows):
        cells = table.rows[i + 1].cells
        for j, val in enumerate(r):
            p = cells[j].paragraphs[0]
            p.clear()
            _append_markdown_inline(p, val)
            _apply_table_alignment(p, aligns[j] if j < len(aligns) else 'left')
        if i % 2 == 1:  # shade every second data row
            for cell in cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F2F2F2')
                tcPr.append(shd)


def _render_markdown_minimal(doc: Document, md: str) -> None:
    """Render a small subset of Markdown (headings/bullets/paragraphs + tables)."""
    if not md:
        return
    lines = md.splitlines()
    para_buf: List[str] = []

    def flush_para():
        nonlocal para_buf
        if para_buf:
            p = doc.add_paragraph()
            _append_markdown_inline(p, " ".join(para_buf).strip())
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(3)
            para_buf = []

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        if not line.strip():
            flush_para()
            i += 1
            continue
        # Table detection
        spec, end_idx = _detect_table(lines, i)
        if spec:
            flush_para()
            _render_table(doc, spec)
            i = end_idx
            continue
        # Headings: support # to ######
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            flush_para()
            hashes, text = m.groups()
            level = len(hashes)
            _add_heading(doc, text.strip(), level=level)
            i += 1
            continue
        if line.lstrip().startswith(("- ", "* ")):
            flush_para()
            bullet = line.lstrip()[2:].strip()
            bp = doc.add_paragraph(style="List Bullet")
            _append_markdown_inline(bp, bullet)
            bp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            bp.paragraph_format.space_after = Pt(0)
            i += 1
            continue
        para_buf.append(line.strip())
        i += 1
    flush_para()
    # Ensure justification where missing (headings excluded)
    for p in doc.paragraphs:
        if p.style and p.style.name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5", "Heading 6"):
            continue
        if p.paragraph_format.alignment is None:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if p.paragraph_format.space_after is None:
            p.paragraph_format.space_after = Pt(3)


def _add_hyperlink(paragraph, url: str, text: str):
    """Add a clickable hyperlink to a paragraph (python-docx low-level)."""
    # Create the relationship id
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed attributes
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Style as Hyperlink
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)

    # Create a w:t element with the link text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _append_markdown_inline(paragraph, text: str):
    """Append inline Markdown (bold, italics, links) as Word runs/hyperlinks."""
    if not text:
        return
    pattern = re.compile(r"(\[[^\]]+\]\([^\)]+\)|\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_)")
    pos = 0
    for m in pattern.finditer(text):
        # preceding plain text
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith('['):
            # link [text](url)
            inner = token[1:token.index(']')]
            url = token[token.index('(') + 1:-1]
            _add_hyperlink(paragraph, url, inner)
        elif token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('__') and token.endswith('__'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('_') and token.endswith('_'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = m.end()
    # trailing text
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_images(doc: Document, images: Iterable[Path]) -> None:
    imgs = list(images or [])
    if not imgs:
        return
    _add_heading(doc, "Figures", level=2)
    for pth in imgs:
        _add_heading(doc, pth.name, level=3)
        try:
            doc.add_picture(str(pth), width=Inches(6.0))
        except Exception:
            # Fallback: mention path if picture fails
            doc.add_paragraph(f"[Image not embedded: {pth}]")
        # Add a small spacer
        sp = doc.add_paragraph("")
        sp.paragraph_format.space_after = Pt(6)


def build_docx_report(
    *,
    title: str,
    meta: dict,
    final_markdown: Optional[str],
    images: Iterable[Path],
    ai_summary: Optional[str],
) -> bytes:
    """Create a .docx report and return it as bytes."""
    doc = Document()
    # Page layout: comfortable margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base font tuning (keeps theme, improves readability)
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
    except Exception:
        pass

    # Title
    _add_heading(doc, title, level=0)

    # Meta overview
    _add_heading(doc, "Overview", level=1)
    for key in ("Question", "Location", "Settings"):
        if key in meta and meta[key]:
            _add_kv(doc, key, str(meta[key]))

    # Final output (markdown)
    if final_markdown:
        _add_heading(doc, "Final Output", level=1)
        _render_markdown_minimal(doc, final_markdown)

    # Images
    _add_images(doc, images)

    # AI Summary
    if ai_summary:
        _add_heading(doc, "AI Summary (process only)", level=1)
        # Preserve bullets if present
        for line in ai_summary.splitlines():
            if line.lstrip().startswith(("- ", "* ")):
                p = doc.add_paragraph(style="List Bullet")
                _append_markdown_inline(p, line.lstrip()[2:].strip())
            else:
                p = doc.add_paragraph()
                _append_markdown_inline(p, line)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(3)

    # Serialize to bytes
    from io import BytesIO

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
